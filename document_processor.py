"""
document_processor.py - Extraktion und *Structural Chunking* fuer UNECE-Dokumente.

Warum kein naives "alle 1000 Zeichen schneiden"?
------------------------------------------------
UNECE-Regelungen sind streng hierarchisch aufgebaut:

    Regulation No. 155
      +- 1.      Scope
      +- 5.      Specifications
      |    +- 5.1.    ...
      |         +- 5.1.2.  ...
      +- Annex 3 - Model of a communication
           +- Appendix 1

Eine Antwort ist juristisch nur brauchbar, wenn die exakte Fundstelle
("Annex 3, para. 5.1.2") mitgeliefert wird. Dieses Modul erkennt daher die
Struktur *vor* dem Chunking und haengt sie als Metadaten an jeden Chunk:

    annex, appendix, paragraph, section, section_path, page_start, page_end,
    regulation, doc_title, source, doc_id, chunk_id, ...

Zusaetzlich wird jedem Chunk-Text ein kompakter Struktur-Header vorangestellt
("[UN R155 | Annex 3 | Para. 5.1.2]"), damit Embedding-Modell *und*
Cross-Encoder-Reranker den Kontext der Passage mitbewerten koennen.

Unterstuetzt: PDF (PyMuPDF/fitz) und DOCX (python-docx).
Das Modul ist bewusst frei von Streamlit-/LangChain-Abhaengigkeiten und laesst
sich damit auch als reine Bibliothek oder im Test nutzen.
"""

from __future__ import annotations

import hashlib
import logging
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Iterator, Sequence

import config

logger = logging.getLogger(__name__)

# PyMuPDF und python-docx werden lazy importiert, damit ein fehlendes Paket
# nur den jeweiligen Dateityp blockiert und nicht das gesamte Modul.
try:  # pragma: no cover - Importpfad haengt von der Installation ab
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None  # type: ignore[assignment]

try:  # pragma: no cover
    import docx  # python-docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore[assignment]


# --------------------------------------------------------------------------- #
# Feature-Schalter
# --------------------------------------------------------------------------- #
#: Tabellen aus PDFs als Markdown extrahieren (etwas langsamer, aber UNECE-
#: Grenzwerte stehen fast immer in Tabellen).
EXTRACT_TABLES: bool = config.env_bool("REG_SEARCH_EXTRACT_TABLES", True)

#: Struktur-Header in den einzubettenden Text schreiben.
PREPEND_STRUCTURE_HEADER: bool = config.env_bool("REG_SEARCH_STRUCTURE_HEADER", True)

#: Kuerzere Textfragmente werden gar nicht erst indiziert.
MIN_USEFUL_CHARS: int = config.env_int("REG_SEARCH_MIN_USEFUL_CHARS", 25)


class DocumentProcessingError(RuntimeError):
    """Fehler beim Einlesen oder Zerlegen eines Dokuments."""


# --------------------------------------------------------------------------- #
# Regulaere Ausdruecke fuer die UNECE-Hierarchie
# --------------------------------------------------------------------------- #
# "Annex 3", "ANNEX 4 - Test procedure", "Anhang II"
RE_ANNEX = re.compile(
    r"^(?:ANNEX|Annex|ANHANG|Anhang)\s+"
    r"(\d{1,2}[A-Za-z]?|[IVXLCDM]{1,6})"
    r"\s*[-–—:.]?\s*(.*)$"
)

# "Appendix 1", "APPENDIX 2 to Annex 3", "Anlage 1"
RE_APPENDIX = re.compile(
    r"^(?:APPENDIX|Appendix|ANLAGE|Anlage)\s+"
    r"(\d{1,2}[A-Za-z]?|[IVXLCDM]{1,6})"
    r"\s*[-–—:.]?\s*(.*)$"
)

# "6.1.2.  Der Hersteller ..." / "5.  Specifications"
# Max. 2 Ziffern in der ersten Ebene verhindert Treffer auf Jahreszahlen.
RE_PARAGRAPH = re.compile(r"^(\d{1,2}(?:\.\d{1,3}){0,6})\.?\s+(?=\S)")

# Reine Nummernzeile ("6.1.2." oder "10." als eigene Zeile, Text folgt darunter).
# Bei einstufigen Nummern ist der Punkt Pflicht - sonst wuerde jede Seitenzahl
# als Abschnitt gelten.
RE_PARAGRAPH_ONLY = re.compile(r"^(\d{1,2}(?:\.\d{1,3})*)\.$|^(\d{1,2}(?:\.\d{1,3})+)$")

# Benannte Abschnitte ohne Nummer (Grossbuchstaben-Ueberschriften)
RE_SECTION_HEADING = re.compile(
    r"^(SCOPE|DEFINITIONS|APPLICATION FOR APPROVAL|APPROVAL|SPECIFICATIONS|"
    r"GENERAL (?:REQUIREMENTS|SPECIFICATIONS|PROVISIONS)|REQUIREMENTS|"
    r"MODIFICATION[S]? (?:AND EXTENSION )?OF .*|CONFORMITY OF PRODUCTION|"
    r"PENALTIES FOR NON-?CONFORMITY .*|PRODUCTION DEFINITIVELY DISCONTINUED|"
    r"NAMES AND ADDRESSES .*|TRANSITIONAL PROVISIONS|MARKINGS?|"
    r"CONTENTS|TABLE OF CONTENTS|INHALT(?:SVERZEICHNIS)?)\s*$"
)

# Fundstellen im Dokumentenkopf
RE_REGULATION_NO = re.compile(
    r"(?:UN\s+)?Regulation\s+No\.?\s*([0-9]{1,3})(?:\s*[-–—]\s*)?", re.IGNORECASE
)
RE_GTR_NO = re.compile(
    r"Global\s+Technical\s+Regulation\s+No\.?\s*([0-9]{1,3})", re.IGNORECASE
)
RE_UN_DOC_SYMBOL = re.compile(
    r"\b(?:E/ECE/[A-Za-z0-9./()-]+|ECE/TRANS/WP\.29[A-Za-z0-9./()-]*)"
)
RE_SERIES = re.compile(
    r"\b(\d{2})\s*series\s+of\s+amendments", re.IGNORECASE
)
# Regelungsnummer aus dem Dateinamen: "R155", "ECE-R79", "UN-R13H"
RE_REG_IN_FILENAME = re.compile(r"\bR[\s_-]?(\d{1,3}[A-Za-z]?)\b", re.IGNORECASE)

# Inhaltsverzeichnis-Zeilen ("5.1  Anforderungen ......... 12")
RE_TOC_LINE = re.compile(r"\.{4,}\s*\d{1,4}\s*$")
# Variante ohne Fuehrungspunkte, wie sie DOCX-Verzeichnisse erzeugen:
# "8<TAB>Reference fuels<TAB><TAB>45". Bewusst eng gefasst - der Titel darf
# keinen Satzpunkt enthalten, damit echte Fliesstextzeilen nicht verschwinden.
RE_TOC_TABBED = re.compile(
    r"^\d{1,2}(?:\.\d{1,3})*\.?[\t ]+[A-Z][^.]{4,90}?[\t ]{2,}\d{1,3}\s*$"
)
# Reine Seitenzahl
RE_PAGE_NUMBER_ONLY = re.compile(r"^[-–—\s]*\d{1,4}[-–—\s]*$")
# Die kanonische Titelformel einer UN-Regelung. Sie steht auf dem Deckblatt oft
# erst *hinter* der Aenderungshistorie, also nicht direkt bei "Regulation No. X".
RE_TITLE_HINT = re.compile(
    r"^(?:Uniform\s+(?:provisions|technical\s+prescriptions)|"
    r"Prescriptions\s+uniformes|Einheitliche\s+(?:Bedingungen|Vorschriften))\b",
    re.IGNORECASE,
)
# Zeilen der Deckblatt-Historie, die nie ein Titel sind
RE_COVER_NOISE = re.compile(
    r"^(?:Supplement|Addendum|Revision|Amendment|Corrigendum|Agreement|"
    r"Concerning\s+the\s+Adoption|Incorporating|Original\s*:)\b",
    re.IGNORECASE,
)
# Datumszeile auf Deckblaettern ("21 August 2013", "1 January 2020")
RE_DATE_LINE = re.compile(
    r"^\d{1,2}\.?\s+(January|February|March|April|May|June|July|August|September|"
    r"October|November|December|Januar|Februar|M(?:ä|ae)rz|Mai|Juni|Juli|Oktober|"
    r"Dezember)\s+\d{4}\.?$",
    re.IGNORECASE,
)


# --------------------------------------------------------------------------- #
# Datenklassen
# --------------------------------------------------------------------------- #
@dataclass
class Chunk:
    """Ein einbettungsfertiges Textstueck inklusive Struktur-Metadaten."""

    text: str
    metadata: dict[str, Any]

    @property
    def chunk_id(self) -> str:
        return str(self.metadata.get("chunk_id", ""))


@dataclass
class ProcessedDocument:
    """Ergebnis der Verarbeitung einer Datei."""

    doc_id: str
    path: Path
    metadata: dict[str, Any]
    chunks: list[Chunk] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def n_chunks(self) -> int:
        return len(self.chunks)


@dataclass
class _Line:
    """Eine Textzeile mit Seitenbezug."""

    text: str
    page: int  # 1-basiert


@dataclass
class _Unit:
    """Eine zusammenhaengende Struktureinheit (z. B. ein Paragraph)."""

    lines: list[str] = field(default_factory=list)
    page_start: int = 1
    page_end: int = 1
    annex: str = ""
    appendix: str = ""
    paragraph: str = ""
    heading: str = ""
    kind: str = "body"  # body | paragraph | annex | appendix | heading | table

    @property
    def text(self) -> str:
        return _join_wrapped_lines(self.lines)

    @property
    def is_empty(self) -> bool:
        return not self.text.strip()


# --------------------------------------------------------------------------- #
# Hilfsfunktionen: Text
# --------------------------------------------------------------------------- #
def _normalise(text: str) -> str:
    """Vereinheitlicht Sonderzeichen, die in UNECE-PDFs haeufig vorkommen."""
    replacements = {
        " ": " ",  # geschuetztes Leerzeichen
        "ﬁ": "fi",
        "ﬂ": "fl",
        "‘": "'",
        "’": "'",
        "“": '"',
        "”": '"',
        "−": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _join_wrapped_lines(lines: Sequence[str]) -> str:
    """
    Fuegt umgebrochene PDF-Zeilen wieder zu Absaetzen zusammen.

    - Zeile endet auf "-"           -> Trennung aufheben ("Fahr-\\nzeug" -> "Fahrzeug")
    - Zeile endet auf .:;!? oder ist eine Tabellen-/Listenzeile -> Zeilenumbruch
    - sonst                         -> mit Leerzeichen verbinden
    """
    out: list[str] = []
    for raw in lines:
        line = raw.strip()
        if not line:
            if out and out[-1] != "":
                out.append("")
            continue
        if not out or out[-1] == "":
            out.append(line)
            continue
        prev = out[-1]
        if prev.endswith("-") and not prev.endswith((" -", "--")):
            out[-1] = prev[:-1] + line
        elif prev.endswith((".", ":", ";", "!", "?", "|")) or line.startswith("|"):
            out.append(line)
        else:
            out[-1] = f"{prev} {line}"
    return "\n".join(out).strip()


def _is_noise(line: str) -> bool:
    """Erkennt Seitenzahlen und Inhaltsverzeichnis-Zeilen."""
    stripped = line.strip()
    if not stripped:
        return False
    return bool(
        RE_PAGE_NUMBER_ONLY.match(stripped)
        or RE_TOC_LINE.search(stripped)
        or RE_TOC_TABBED.match(line.rstrip())
    )


#: Satzzeichen, die das Ende eines Absatzes markieren.
_BLOCK_END_CHARS = (".", ":", ";", "!", "?", "|", ")")


def _starts_new_block(previous_line: str) -> bool:
    """
    Beginnt hier ein neuer Absatz - oder laeuft nur ein Zeilenumbruch weiter?

    Entscheidend fuer die Unterscheidung von echten Ueberschriften und
    Querverweisen im Fliesstext.
    """
    previous = previous_line.strip()
    return not previous or previous.endswith(_BLOCK_END_CHARS)


def _is_heading_candidate(line: str, match: re.Match[str], previous_line: str) -> bool:
    """
    Trennt echte Annex-/Appendix-Ueberschriften von Querverweisen im Fliesstext.

    Gegenbeispiel aus der Praxis::

        3.2.1. a description of the vehicle type with regard to the items
        specified in Annex 1 to this Regulation;

    Zeile 2 beginnt mit "Annex 1", ist aber die Fortsetzung eines Satzes.
    Wuerde sie als Ueberschrift gewertet, bekaemen alle folgenden Chunks eine
    falsche Fundstelle.

    Entscheidend ist, was *hinter* dem Marker steht - nicht die Vorzeile:
    In UN R85 folgt die Ueberschrift "Annex 3a" direkt auf eine Formularzeile
    ("Circulation: with / without"), die kein Satzzeichen hat. Eine Pruefung
    auf die Vorzeile allein verwirft diese echte Ueberschrift faelschlich.
    """
    stripped = line.rstrip()
    if len(stripped) > 100 or stripped.endswith((",", ";")):
        return False

    title = (match.group(2) or "").strip()
    if title:
        # "Annex 1 to this Regulation" -> Verweis; "Annex 6 Method for ..." -> Titel
        # Ein abschliessender Punkt verraet einen Satz, keine Ueberschrift.
        return not (title[0].islower() or title.endswith("."))

    # Nur der Marker: "Annex 3a" ist eine Ueberschrift, "... in Annex 5." nicht.
    return not stripped.endswith(".") or _starts_new_block(previous_line)


def _paragraph_key(number: str) -> tuple[int, ...]:
    """Zerlegt "5.4.2.2" in ``(5, 4, 2, 2)`` - damit werden Nummern vergleichbar."""
    try:
        return tuple(int(part) for part in number.rstrip(".").split("."))
    except ValueError:  # pragma: no cover - vom Regex ausgeschlossen
        return ()


def _is_plausible_number(key: tuple[int, ...]) -> bool:
    """
    Sieht diese Nummer nach einem Abschnitt aus - oder nach einem Messwert?

    UNECE-Nummerierungen beginnen auf jeder Ebene bei 1. Eine Null als
    Unterebene ("7.0") oder als Vorsatz einer mehrstelligen Nummer ("0.3")
    kommt darin nicht vor - solche Zeilen sind Grenzwerte oder Fragmente einer
    vom PDF-Extraktor aufgebrochenen Formel.

    Gegenbeispiele aus UN R85, Annex 5::

        fa =  (99/Ps)^0.7          ->  Zeile "7.0"
        ... a constant value of fm equal to
        0.3 (fm =  0.3) will be taken.

    Beide eroeffneten frueher einen Paragraphen. Beim zweiten brach der Satz
    genau vor dem Grenzwert ab: die Antwort nannte einen falschen Wert, und die
    Fundstelle dazu sah trotzdem korrekt aus. Das ist die gefaehrlichste
    Fehlerart dieses Werkzeugs - im Zweifel deshalb kein neuer Abschnitt.

    Eine einstellige Null ("0." als Einleitung) bleibt erlaubt; nur in
    Kombination mit weiteren Ebenen ist sie ein sicheres Artefakt.
    """
    if not key:
        return False
    if len(key) > 1 and (key[0] == 0 or any(part == 0 for part in key[1:])):
        return False
    return True


RE_HAS_WORD = re.compile(r"[A-Za-zÀ-ɏ]{2,}")


def _has_text_below(lines: Sequence[_Line], position: int) -> bool:
    """
    Folgt auf eine reine Nummernzeile echter Text - oder nur eine Formel?

    In UNECE-PDFs steht die Ueberschrift unmittelbar unter ihrer Nummer::

        5.4.2.2.
        Engine factor fm

    Aufgebrochene Formeln sehen anders aus: PyMuPDF liefert den Exponenten von
    ``(99/Ps)^1.5 * (T/298)^0.7`` als eigene Zeilen "5.1" und "7.0", gefolgt von
    Klammer- und Operatorglyphen. Ohne diese Pruefung eroeffnet "5.1" einen
    Paragraphen; beim Zusammenfassen kleiner Einheiten wandert die Fundstelle
    dann auf den gemeinsamen Praefix "5" und wird unbrauchbar grob.
    """
    for candidate in lines[position + 1 : position + 4]:
        text = candidate.text.strip()
        if text:
            return bool(RE_HAS_WORD.search(text))
    return False


def _is_paragraph_start(
    line: str, match: re.Match[str], previous_line: str, last_top_level: int
) -> bool:
    """
    Prueft, ob eine Nummer am Zeilenanfang wirklich einen Paragraphen eroeffnet.

    Ohne diese Pruefung entstehen aus Messwerten und Titelfragmenten falsche
    Fundstellen - in UN R85 etwa aus der Annex-6-Ueberschrift "... and the
    maximum 30 minutes power ..." der Paragraph "30", oder aus dem Deckblatt-
    Datum "21 August 2013" der Paragraph "21". Eine falsche Fundstelle ist
    schlimmer als eine grobe: im Zweifel bleibt der Text beim vorherigen
    Paragraphen.

    Args:
        last_top_level: hoechste bisher gesehene Abschnittsnummer der aktuellen
            Ebene (wird bei jedem Annex zurueckgesetzt). UNECE-Abschnitte laufen
            aufsteigend, ein Sprung von 1 auf 30 ist praktisch immer ein Artefakt.
    """
    number = match.group(1)
    rest = line[match.end() :].lstrip()

    try:
        top_level = int(number.split(".")[0])
    except ValueError:  # pragma: no cover - vom Regex ausgeschlossen
        return False

    # Plausibilitaet der Nummernfolge (kleine Spruenge fuer gestrichene
    # Abschnitte erlaubt)
    if top_level > last_top_level + 5:
        return False

    if "." not in number:
        # Einstufige Abschnitte tragen in UNECE-Texten immer eine Ueberschrift
        # in Grossschreibung ("3. Application for approval") - "30 minutes
        # power of electric drive trains" faellt damit heraus.
        if not (rest[:1].isupper() or rest[:1] in {'"', "'", "("}):
            return False
        return True

    if _starts_new_block(previous_line):
        return True
    # Mehrstufige Nummern ("3.2.1.") sind ein starkes Signal - dann muss aber
    # der Folgetext wie ein Satzanfang aussehen. Laeuft ein Satz aus der
    # Vorzeile weiter, ist eine oeffnende Klammer fast immer eine Formel oder
    # Einheit ("... equal to" / "0.3 (fm = 0.3) will be taken") - hier zaehlt
    # deshalb nur ein echter Satzanfang.
    return bool(rest[:1].isupper() or rest[:1] in {'"', "'"})


def _is_title_like(text: str, max_len: int = 90) -> bool:
    """
    Sieht dieser Text nach einer Abschnittsueberschrift aus?

    Ueberschriften in UNECE-Texten sind kurz, beginnen gross und enden ohne
    Punkt ("Application for approval"). Fliesstext ("The application shall be
    submitted ...") faellt damit heraus.
    """
    stripped = text.strip()
    if not (0 < len(stripped) <= max_len):
        return False
    if stripped.endswith((".", ";", ":")) or stripped.startswith("|"):
        return False
    return bool(stripped[:1].isupper())


def _looks_like_heading(line: str) -> bool:
    """Grossbuchstaben-Ueberschrift ohne Satzzeichen (z. B. 'SCOPE')."""
    stripped = line.strip()
    if not (3 <= len(stripped) <= 90):
        return False
    if RE_SECTION_HEADING.match(stripped):
        return True
    letters = [c for c in stripped if c.isalpha()]
    if len(letters) < 3:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio > 0.9 and not stripped.endswith((".", ",", ";"))


def sanitize_metadata(metadata: dict[str, Any]) -> dict[str, Any]:
    """
    ChromaDB akzeptiert nur skalare Metadaten (str, int, float, bool).

    ``None`` wird verworfen, Listen werden zu kommaseparierten Strings.
    """
    clean: dict[str, Any] = {}
    for key, value in metadata.items():
        if value is None or value == "":
            continue
        if isinstance(value, (str, int, float, bool)):
            clean[key] = value
        elif isinstance(value, (list, tuple, set)):
            joined = ", ".join(str(v) for v in value if v not in (None, ""))
            if joined:
                clean[key] = joined
        else:
            clean[key] = str(value)
    return clean


def compute_doc_id(path: Path) -> str:
    """Inhaltsbasierte ID (SHA-1 der Datei) - erkennt Duplikate zuverlaessig."""
    digest = hashlib.sha1()
    with open(path, "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()[:16]


# --------------------------------------------------------------------------- #
# PDF-Extraktion
# --------------------------------------------------------------------------- #
def _rows_to_markdown(rows: Sequence[Sequence[Any]]) -> str:
    """Wandelt eine extrahierte Tabelle in eine kompakte Markdown-Tabelle."""
    cleaned: list[list[str]] = []
    for row in rows:
        cells = [(str(c).replace("\n", " ").strip() if c is not None else "") for c in row]
        if any(cells):
            cleaned.append(cells)
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    cleaned = [r + [""] * (width - len(r)) for r in cleaned]
    header, *body = cleaned
    lines = ["| " + " | ".join(header) + " |", "|" + "|".join([" --- "] * width) + "|"]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


def _extract_pdf_tables(page: Any) -> list[tuple[Any, str]]:
    """Findet Tabellen einer Seite und liefert (Rect, Markdown)-Paare."""
    if not EXTRACT_TABLES:
        return []
    try:
        finder = page.find_tables()
    except Exception as exc:  # pragma: no cover - abhaengig von PyMuPDF-Version
        logger.debug("Tabellensuche auf Seite %s fehlgeschlagen: %s", page.number, exc)
        return []

    tables: list[tuple[Any, str]] = []
    for table in getattr(finder, "tables", []):
        try:
            markdown = _rows_to_markdown(table.extract())
            if markdown:
                tables.append((fitz.Rect(table.bbox), markdown))
        except Exception as exc:  # pragma: no cover
            logger.debug("Tabelle konnte nicht extrahiert werden: %s", exc)
    return tables


def _pdf_lines(document: Any) -> tuple[list[_Line], list[str]]:
    """
    Liest ein PDF seitenweise in eine flache Zeilenliste.

    Tabellen werden - sofern erkannt - als Markdown-Bloecke an ihrer Position im
    Lesefluss eingefuegt; die zugehoerigen Textbloecke werden uebersprungen,
    damit Zelleninhalte nicht doppelt auftauchen.
    """
    lines: list[_Line] = []
    warnings: list[str] = []
    empty_pages = 0

    for page_index in range(document.page_count):
        page = document.load_page(page_index)
        page_no = page_index + 1

        tables = _extract_pdf_tables(page)
        table_rects = [rect for rect, _ in tables]

        items: list[tuple[float, str]] = []
        try:
            blocks = page.get_text("blocks", sort=True)
        except Exception as exc:  # pragma: no cover
            warnings.append(f"Seite {page_no} konnte nicht gelesen werden: {exc}")
            continue

        for block in blocks:
            x0, y0, x1, y1, text = block[0], block[1], block[2], block[3], block[4]
            block_type = block[6] if len(block) > 6 else 0
            if block_type != 0 or not str(text).strip():
                continue
            block_rect = fitz.Rect(x0, y0, x1, y1)
            # Bloecke innerhalb einer erkannten Tabelle auslassen
            if any(rect.intersects(block_rect) and rect & block_rect for rect in table_rects):
                overlap = max(
                    ((rect & block_rect).get_area() / max(block_rect.get_area(), 1e-6))
                    for rect in table_rects
                    if rect.intersects(block_rect)
                )
                if overlap > 0.5:
                    continue
            items.append((float(y0), str(text)))

        for rect, markdown in tables:
            items.append((float(rect.y0), "\n" + markdown + "\n"))

        items.sort(key=lambda item: item[0])
        page_text = "\n".join(text for _, text in items)
        if not page_text.strip():
            empty_pages += 1
        for line in _normalise(page_text).split("\n"):
            lines.append(_Line(text=line.rstrip(), page=page_no))

    if document.page_count and empty_pages / document.page_count > 0.5:
        warnings.append(
            "Mehr als die Haelfte der Seiten enthaelt keinen Textlayer - "
            "vermutlich ein gescanntes PDF. Bitte vorher OCR anwenden "
            "(z. B. ocrmypdf), sonst bleibt der Inhalt unsuchbar."
        )
    return lines, warnings


def _drop_repeating_lines(lines: list[_Line], n_pages: int) -> list[_Line]:
    """
    Entfernt Kopf-/Fusszeilen (z. B. 'ECE/TRANS/WP.29/2021/xx').

    Kurze Zeilen, die auf mindestens ``HEADER_FOOTER_RATIO`` der Seiten
    vorkommen, gelten als Boilerplate.
    """
    if n_pages < 4:
        return lines

    pages_per_line: dict[str, set[int]] = {}
    for line in lines:
        key = line.text.strip()
        if 0 < len(key) <= 120:
            pages_per_line.setdefault(key, set()).add(line.page)

    threshold = max(3, int(n_pages * config.HEADER_FOOTER_RATIO))
    boilerplate = {key for key, pages in pages_per_line.items() if len(pages) >= threshold}
    if boilerplate:
        logger.debug("Boilerplate entfernt: %s", sorted(boilerplate)[:5])
    return [line for line in lines if line.text.strip() not in boilerplate]


# --------------------------------------------------------------------------- #
# DOCX-Extraktion
# --------------------------------------------------------------------------- #
def _iter_docx_blocks(document: Any) -> Iterator[tuple[str, Any]]:
    """Iteriert Absaetze und Tabellen eines DOCX in Dokumentreihenfolge."""
    from docx.table import Table  # lokal: nur bei DOCX benoetigt
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield "p", Paragraph(child, document)
        elif tag == "tbl":
            yield "table", Table(child, document)


def _docx_lines(path: Path) -> tuple[list[_Line], list[str]]:
    """Liest ein DOCX in dieselbe flache Zeilenliste wie ``_pdf_lines``."""
    if docx is None:  # pragma: no cover
        raise DocumentProcessingError(
            "python-docx ist nicht installiert - bitte 'pip install python-docx'."
        )
    document = docx.Document(str(path))
    lines: list[_Line] = []
    warnings: list[str] = []

    for kind, block in _iter_docx_blocks(document):
        if kind == "p":
            text = _normalise(block.text).strip()
            if not text:
                lines.append(_Line("", 0))
                continue
            style = (getattr(block.style, "name", "") or "").lower()
            # Word-Ueberschriften ohne eigene Nummerierung als Heading markieren
            if style.startswith("heading") and not RE_PARAGRAPH.match(text):
                lines.append(_Line("", 0))
                lines.append(_Line(text, 0))
                lines.append(_Line("", 0))
            else:
                lines.append(_Line(text, 0))
        else:
            rows = [[cell.text for cell in row.cells] for row in block.rows]
            markdown = _rows_to_markdown(rows)
            if markdown:
                lines.append(_Line("", 0))
                for row_line in markdown.split("\n"):
                    lines.append(_Line(row_line, 0))
                lines.append(_Line("", 0))

    if not any(line.text.strip() for line in lines):
        warnings.append("Das DOCX enthaelt keinen extrahierbaren Text.")
    return lines, warnings


# --------------------------------------------------------------------------- #
# Structural Chunking
# --------------------------------------------------------------------------- #
def _detect_document_meta(lines: Sequence[_Line], path: Path) -> dict[str, Any]:
    """
    Ermittelt Regelungsnummer, Serie, UN-Dokumentsymbol und Titel.

    Gesucht wird primaer im Dokumentenkopf (erste ~120 Zeilen bzw. 3 Seiten),
    ersatzweise im Dateinamen.
    """
    head_lines = [line.text for line in lines if line.page <= 3][:200] or [
        line.text for line in lines[:200]
    ]
    head = "\n".join(head_lines)

    meta: dict[str, Any] = {}

    match = RE_REGULATION_NO.search(head)
    gtr = RE_GTR_NO.search(head)
    if match:
        meta["regulation_no"] = match.group(1)
        meta["regulation"] = f"UN Regulation No. {match.group(1)}"
    elif gtr:
        meta["regulation_no"] = f"GTR {gtr.group(1)}"
        meta["regulation"] = f"UN Global Technical Regulation No. {gtr.group(1)}"
    else:
        filename_match = RE_REG_IN_FILENAME.search(path.stem)
        if filename_match:
            meta["regulation_no"] = filename_match.group(1)
            meta["regulation"] = f"UN Regulation No. {filename_match.group(1)}"

    symbol = RE_UN_DOC_SYMBOL.search(head)
    if symbol:
        meta["un_symbol"] = symbol.group(0).rstrip(".,;")

    series = RE_SERIES.search(head)
    if series:
        meta["amendment_series"] = f"{series.group(1)} series of amendments"

    # Titel: bevorzugt die beschreibende Zeile direkt hinter "Regulation No. X"
    # ("Uniform provisions concerning the approval of ..."). Deckblaetter
    # enthalten sonst Datumszeilen ("21 August 2013"), Organisationsnamen
    # ("UNITED NATIONS") oder - bei DOCX - Reste von Tabellenrahmen.
    def usable(candidate: str) -> bool:
        candidate = candidate.strip()
        return (
            15 <= len(candidate) <= 300
            and not candidate.startswith("|")
            and not candidate.endswith(":")  # "Incorporating all valid text up to:"
            and not RE_UN_DOC_SYMBOL.search(candidate)
            and not RE_PAGE_NUMBER_ONLY.match(candidate)
            and not RE_DATE_LINE.match(candidate)
            and not RE_COVER_NOISE.match(candidate)
            and not candidate.isupper()
            and not candidate.lower().startswith(("page", "seite"))
            and any(character.isalpha() for character in candidate)
        )

    def collect(start: int) -> str:
        """Setzt einen ueber mehrere Zeilen umbrochenen Titel zusammen."""
        parts: list[str] = []
        for follow in head_lines[start : start + 4]:
            if not usable(follow):
                if parts:
                    break
                continue
            parts.append(follow.strip())
            joined = " ".join(parts)
            if joined.endswith(".") or len(joined) > 240:
                break
        return " ".join(parts).rstrip(".")

    # 1. Vorzug: die kanonische Titelformel ("Uniform provisions concerning ...")
    title = ""
    for index, text in enumerate(head_lines):
        if RE_TITLE_HINT.match(text.strip()):
            title = collect(index)
            break

    # 2. Ersatz: die Zeile direkt hinter "Regulation No. X" (andere Layouts)
    if not title:
        for index, text in enumerate(head_lines):
            if RE_REGULATION_NO.search(text) or RE_GTR_NO.search(text):
                title = collect(index + 1)
                if title:
                    break

    # 3. Notnagel: die erste brauchbare Zeile ueberhaupt
    if not title:
        title = collect(0)

    # 240 Zeichen schnitten UNECE-Titel mitten im Sachteil ab: bei UN R85
    # endete der gespeicherte Titel bei "... the maximum 30 minutes" und der
    # daraus abgeleitete Kurztitel damit ebenfalls. Volle Titel sind ~300.
    meta["doc_title"] = _shorten(title, 400) if title else path.stem
    return meta


def _parse_units(lines: Sequence[_Line]) -> list[_Unit]:
    """
    Zerlegt die Zeilenliste anhand der UNECE-Hierarchie in Struktureinheiten.

    Zustandsautomat: ``annex``/``appendix``/``paragraph`` bleiben so lange
    gueltig, bis ein neuer Marker auftaucht - dadurch erbt auch Fliesstext
    ohne eigene Nummer die korrekte Fundstelle.
    """
    units: list[_Unit] = []
    current_annex = ""
    current_appendix = ""
    current_paragraph = ""
    current_heading = ""

    unit = _Unit(page_start=lines[0].page if lines else 1, page_end=1)

    def flush() -> None:
        nonlocal unit
        if not unit.is_empty:
            units.append(unit)
        unit = _Unit(
            page_start=1,
            page_end=1,
            annex=current_annex,
            appendix=current_appendix,
            paragraph=current_paragraph,
            heading=current_heading,
        )

    pending_title = False  # naechste Zeile ist die Ueberschrift des Abschnitts
    previous_text = ""  # letzte verarbeitete Zeile (fuer die Umbruch-Erkennung)
    last_top_level = 0  # hoechste Abschnittsnummer der aktuellen Ebene

    for position, line in enumerate(lines):
        text = line.text.strip()

        # Vorzeile sichern und sofort fortschreiben - so bleibt der Zustand
        # auch ueber die vielen 'continue'-Zweige hinweg korrekt.
        previous = previous_text
        previous_text = text

        if _is_noise(text):
            previous_text = ""  # Seitenzahl trennt optisch -> neuer Block
            continue

        # Laufende Kopfzeile direkt ueber der gleichlautenden Ueberschrift
        # ("Annex 1" als Seitenkopf *und* als Titel) nicht doppelt uebernehmen.
        if text == previous and (RE_ANNEX.match(text) or RE_APPENDIX.match(text)):
            continue

        # --- Annex ---------------------------------------------------------
        annex_match = RE_ANNEX.match(text)
        if annex_match and _is_heading_candidate(text, annex_match, previous):
            flush()
            current_annex = annex_match.group(1)
            current_appendix = ""
            current_paragraph = ""
            current_heading = (annex_match.group(2) or "").strip()
            unit = _Unit(
                page_start=line.page,
                page_end=line.page,
                annex=current_annex,
                heading=current_heading,
                kind="annex",
                lines=[text],
            )
            pending_title = not current_heading
            last_top_level = 0  # Nummerierung startet im Annex neu
            continue

        # --- Appendix ------------------------------------------------------
        appendix_match = RE_APPENDIX.match(text)
        if appendix_match and _is_heading_candidate(text, appendix_match, previous):
            flush()
            current_appendix = appendix_match.group(1)
            current_paragraph = ""
            current_heading = (appendix_match.group(2) or "").strip()
            unit = _Unit(
                page_start=line.page,
                page_end=line.page,
                annex=current_annex,
                appendix=current_appendix,
                heading=current_heading,
                kind="appendix",
                lines=[text],
            )
            pending_title = not current_heading
            last_top_level = 0
            continue

        # Titelzeile direkt unter "Annex 3" bzw. unter einer Abschnittsnummer
        if pending_title and text:
            pending_title = False
            if _is_title_like(text, max_len=140) and not RE_PARAGRAPH.match(text):
                current_heading = text
                unit.heading = text
                unit.lines.append(text)
                unit.page_end = line.page
                continue

        # --- Nummerierter Paragraph ----------------------------------------
        paragraph_only = RE_PARAGRAPH_ONLY.match(text)
        paragraph_match = RE_PARAGRAPH.match(text)

        # Erst die Nummer bestimmen, dann pruefen: die Plausibilitaet der
        # Nummernfolge gilt fuer beide Schreibweisen. Die reine Nummernzeile
        # ("7.0" als eigene Zeile) lief frueher ungeprueft durch und erzeugte
        # aus Formelfragmenten Phantom-Abschnitte.
        if paragraph_only:
            number = (paragraph_only.group(1) or paragraph_only.group(2)).rstrip(".")
            remainder = ""
            opens_paragraph = _is_plausible_number(
                _paragraph_key(number)
            ) and _has_text_below(lines, position)
        elif paragraph_match:
            number = paragraph_match.group(1).rstrip(".")
            remainder = text[paragraph_match.end() :].strip()
            opens_paragraph = _is_plausible_number(
                _paragraph_key(number)
            ) and _is_paragraph_start(text, paragraph_match, previous, last_top_level)
        else:
            number = ""
            remainder = ""
            opens_paragraph = False

        if opens_paragraph:
            flush()
            current_paragraph = number
            try:
                last_top_level = max(last_top_level, int(number.split(".")[0]))
            except ValueError:  # pragma: no cover
                pass

            # Einstufige Nummern eroeffnen einen benannten Abschnitt
            # ("3. Application for approval"). Ohne das Zuruecksetzen wuerde
            # die Ueberschrift des Deckblatts ("United Nations") bis zum
            # Dokumentende an allen Fundstellen kleben.
            if "." not in number:
                if _is_title_like(remainder):
                    current_heading = remainder
                    pending_title = False
                else:
                    current_heading = ""
                    pending_title = not remainder

            unit = _Unit(
                page_start=line.page,
                page_end=line.page,
                annex=current_annex,
                appendix=current_appendix,
                paragraph=current_paragraph,
                heading=current_heading,
                kind="paragraph",
                lines=[text],
            )
            continue

        # --- Benannte Ueberschrift -----------------------------------------
        if _looks_like_heading(text):
            flush()
            current_heading = text.title() if text.isupper() else text
            current_paragraph = ""
            unit = _Unit(
                page_start=line.page,
                page_end=line.page,
                annex=current_annex,
                appendix=current_appendix,
                heading=current_heading,
                kind="heading",
                lines=[text],
            )
            continue

        # --- Fliesstext ------------------------------------------------------
        if not unit.lines:
            unit.page_start = line.page
        unit.lines.append(text)
        unit.page_end = max(unit.page_end, line.page)

    if not unit.is_empty:
        units.append(unit)
    return units


def _split_text(text: str, size: int, overlap: int) -> list[str]:
    """
    Teilt zu lange Struktureinheiten rekursiv an natuerlichen Grenzen.

    Reihenfolge der Trenner: Absatz -> Zeile -> Satz -> Wort. Das haelt
    Aufzaehlungen und Saetze zusammen, bevor hart geschnitten wird.
    """
    text = text.strip()
    if len(text) <= size:
        return [text] if text else []

    separators = ["\n\n", "\n", ". ", "; ", ", ", " "]
    for separator in separators:
        if separator not in text:
            continue
        parts = text.split(separator)
        chunks: list[str] = []
        buffer = ""
        for part in parts:
            candidate = part if not buffer else f"{buffer}{separator}{part}"
            if len(candidate) <= size:
                buffer = candidate
                continue
            if buffer:
                chunks.append(buffer)
            if len(part) > size:
                chunks.extend(_split_text(part, size, overlap))
                buffer = ""
            else:
                buffer = part
        if buffer:
            chunks.append(buffer)
        if len(chunks) > 1:
            return _apply_overlap(chunks, overlap)

    # Kein Trenner gefunden -> harter Schnitt
    step = max(size - overlap, 1)
    return [text[i : i + size] for i in range(0, len(text), step)]


def _apply_overlap(chunks: list[str], overlap: int) -> list[str]:
    """
    Haengt jedem Chunk den Schwanz seines Vorgaengers als Kontext an.

    Der Schnitt erfolgt an einer Wortgrenze - sonst beginnt der Chunk mit einem
    Fragment ("nnex 3b to this Regulation"), das sowohl das Embedding als auch
    das angezeigte Zitat verunstaltet.
    """
    if overlap <= 0 or len(chunks) < 2:
        return chunks
    result = [chunks[0]]
    for previous, current in zip(chunks, chunks[1:]):
        tail = previous[-overlap:]
        cut = tail.find(" ")
        tail = tail[cut + 1 :].strip() if cut != -1 else ""
        result.append(f"{tail} {current}".strip() if tail else current)
    return result


def _shorten(text: str, limit: int = 48) -> str:
    """Kuerzt an der Wortgrenze, damit Zitate nicht mitten im Wort abbrechen."""
    text = " ".join(str(text).split())
    if len(text) <= limit:
        return text
    cut = text[:limit].rsplit(" ", 1)[0]
    return f"{cut or text[:limit]}..."


def build_citation(metadata: dict[str, Any]) -> str:
    """Baut die menschenlesbare Fundstelle, z. B. 'UN R155 · Annex 3 · Para. 5.1.2 · S. 12'."""
    parts: list[str] = []
    regulation_no = metadata.get("regulation_no")
    if regulation_no:
        parts.append(f"UN R{regulation_no}")
    else:
        title = str(metadata.get("doc_title") or metadata.get("source") or "")
        if title:
            parts.append(_shorten(title))

    if metadata.get("annex"):
        parts.append(f"Annex {metadata['annex']}")
    if metadata.get("appendix"):
        parts.append(f"Appendix {metadata['appendix']}")
    if metadata.get("paragraph"):
        parts.append(f"Para. {metadata['paragraph']}")
    elif metadata.get("heading"):
        parts.append(_shorten(metadata["heading"]))

    # DOCX kennt keine Seiten (page_start == 0) - dann entfaellt die Angabe.
    page_start = metadata.get("page_start")
    page_end = metadata.get("page_end")
    if page_start:
        parts.append(f"S. {page_start}" if page_start == page_end else f"S. {page_start}-{page_end}")
    return " · ".join(parts)


def _structure_header(metadata: dict[str, Any]) -> str:
    """Kompakter Kontext-Header, der mit eingebettet wird."""
    parts: list[str] = []
    if metadata.get("regulation"):
        parts.append(str(metadata["regulation"]))
    if metadata.get("annex"):
        parts.append(f"Annex {metadata['annex']}")
    if metadata.get("appendix"):
        parts.append(f"Appendix {metadata['appendix']}")
    if metadata.get("paragraph"):
        parts.append(f"Para. {metadata['paragraph']}")
    if metadata.get("heading"):
        parts.append(str(metadata["heading"]))
    return "[" + " | ".join(parts) + "]" if parts else ""


def _merge_small_units(units: list[_Unit], min_chars: int, max_chars: int) -> list[_Unit]:
    """
    Fasst zu kleine Einheiten mit der jeweils folgenden zusammen.

    Verhindert Mini-Chunks wie "6.1.2." oder einzeilige Ueberschriften und
    haelt gleichzeitig den Kontext beisammen. Zusammengefasst wird nur
    innerhalb desselben Annex/Appendix.

    Einheiten mit *verschiedenen* Paragraphennummern werden nur noch
    zusammengefasst, wenn die kleinere sonst unter MIN_USEFUL_CHARS
    herausfiele; die Fundstelle wandert dann auf den gemeinsamen
    Oberabschnitt (5.4.2.1.1 + 5.4.2.2 -> "Para. 5.4.2").

    Frueher gewann die Nummer der ersten Einheit: der Chunk trug eine
    Fundstelle, unter der ein Teil seines Textes gar nicht steht. Der haeufige
    Fall "Ueberschrift 5.2 + Unterabschnitte 5.2.1/5.2.2" bleibt unberuehrt -
    der gemeinsame Praefix ist dort 5.2.
    """
    def same_section(a: _Unit, b: _Unit) -> bool:
        """Gleicher Annex *und* gleicher Hauptabschnitt (5.x darf nicht zu 6.x)."""
        if a.annex != b.annex or a.appendix != b.appendix:
            return False
        top_a = a.paragraph.split(".")[0] if a.paragraph else ""
        top_b = b.paragraph.split(".")[0] if b.paragraph else ""
        return not (top_a and top_b) or top_a == top_b

    def common_paragraph(a: _Unit, b: _Unit) -> str | None:
        """Fundstelle, unter der *beide* Texte tatsaechlich stehen."""
        if not a.paragraph or not b.paragraph or a.paragraph == b.paragraph:
            return a.paragraph or b.paragraph
        common: list[str] = []
        for part_a, part_b in zip(a.paragraph.split("."), b.paragraph.split(".")):
            if part_a != part_b:
                break
            common.append(part_a)
        return ".".join(common) if common else None

    merged: list[_Unit] = []
    for unit in units:
        if (
            merged
            and len(merged[-1].text) < min_chars
            and same_section(merged[-1], unit)
            and len(merged[-1].text) + len(unit.text) <= max_chars
        ):
            previous = merged[-1]
            paragraph = common_paragraph(previous, unit)
            # Wandert die Fundstelle dabei auf einen Oberabschnitt, geht
            # Genauigkeit verloren. Das lohnt nur, wenn die kleine Einheit
            # sonst ganz herausfiele (< MIN_USEFUL_CHARS) - andernfalls sind
            # zwei exakt zitierte Chunks mehr wert als ein grob zitierter.
            coarsens = bool(previous.paragraph) and paragraph != previous.paragraph
            if paragraph is not None and (
                not coarsens or len(previous.text) < MIN_USEFUL_CHARS
            ):
                previous.lines.extend(unit.lines)
                previous.page_end = max(previous.page_end, unit.page_end)
                previous.paragraph = paragraph
                if not previous.heading and unit.heading:
                    previous.heading = unit.heading
                continue
        merged.append(unit)
    return merged


def units_to_chunks(
    units: Sequence[_Unit],
    doc_metadata: dict[str, Any],
    chunk_size: int,
    chunk_overlap: int,
) -> list[Chunk]:
    """Wandelt Struktureinheiten in einbettungsfertige Chunks um."""
    chunks: list[Chunk] = []
    doc_id = str(doc_metadata.get("doc_id", ""))

    for unit in units:
        body = unit.text.strip()
        if not body:
            continue

        for part_index, part in enumerate(_split_text(body, chunk_size, chunk_overlap)):
            # Winzige Fragmente (Formularreste, Trennzeichen) tragen nichts zur
            # Suche bei, verwaessern aber die Trefferliste.
            if len(part.strip()) < MIN_USEFUL_CHARS:
                continue
            index = len(chunks)
            metadata: dict[str, Any] = {
                **doc_metadata,
                "annex": unit.annex,
                "appendix": unit.appendix,
                "paragraph": unit.paragraph,
                "heading": unit.heading,
                "section_type": unit.kind,
                "page_start": unit.page_start,
                "page_end": unit.page_end,
                "chunk_index": index,
                "part_index": part_index,
                "chunk_id": f"{doc_id}:{index:05d}",
                "char_count": len(part),
            }
            metadata["section_path"] = " > ".join(
                piece
                for piece in (
                    f"Annex {unit.annex}" if unit.annex else "",
                    f"Appendix {unit.appendix}" if unit.appendix else "",
                    unit.heading or "",
                    f"Para. {unit.paragraph}" if unit.paragraph else "",
                )
                if piece
            )
            metadata["citation"] = build_citation(metadata)

            header = _structure_header(metadata) if PREPEND_STRUCTURE_HEADER else ""
            text = f"{header}\n{part}".strip() if header else part
            chunks.append(Chunk(text=text, metadata=sanitize_metadata(metadata)))
    return chunks


# --------------------------------------------------------------------------- #
# Oeffentliche API
# --------------------------------------------------------------------------- #
def process_document(
    path: str | Path,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> ProcessedDocument:
    """
    Liest eine PDF-/DOCX-Datei und erzeugt struktur-annotierte Chunks.

    Raises:
        DocumentProcessingError: bei nicht unterstuetztem Format, fehlender
            Abhaengigkeit oder unlesbarer Datei.
    """
    path = Path(path)
    chunk_size = chunk_size or config.CHUNK_SIZE
    chunk_overlap = chunk_overlap or config.CHUNK_OVERLAP

    if not path.is_file():
        raise DocumentProcessingError(f"Datei nicht gefunden: {path}")
    suffix = path.suffix.lower()
    if suffix not in config.SUPPORTED_EXTENSIONS:
        raise DocumentProcessingError(
            f"Format '{suffix}' wird nicht unterstuetzt "
            f"(erlaubt: {', '.join(config.SUPPORTED_EXTENSIONS)})."
        )

    doc_id = compute_doc_id(path)
    warnings: list[str] = []
    pdf_meta: dict[str, Any] = {}
    n_pages = 1

    if suffix == ".pdf":
        if fitz is None:  # pragma: no cover
            raise DocumentProcessingError(
                "PyMuPDF ist nicht installiert - bitte 'pip install pymupdf'."
            )
        try:
            with fitz.open(str(path)) as document:
                if document.is_encrypted and not document.authenticate(""):
                    raise DocumentProcessingError(
                        f"'{path.name}' ist passwortgeschuetzt und kann nicht gelesen werden."
                    )
                n_pages = document.page_count
                lines, extraction_warnings = _pdf_lines(document)
                pdf_meta = {
                    key: value
                    for key, value in (document.metadata or {}).items()
                    if key in {"title", "author", "subject", "creationDate"} and value
                }
        except DocumentProcessingError:
            raise
        except Exception as exc:  # pragma: no cover - defekte PDFs
            raise DocumentProcessingError(f"PDF konnte nicht geoeffnet werden: {exc}") from exc
        warnings.extend(extraction_warnings)
        lines = _drop_repeating_lines(lines, n_pages)
    else:
        try:
            lines, extraction_warnings = _docx_lines(path)
        except DocumentProcessingError:
            raise
        except Exception as exc:  # pragma: no cover
            raise DocumentProcessingError(f"DOCX konnte nicht gelesen werden: {exc}") from exc
        warnings.extend(extraction_warnings)

    if not any(line.text.strip() for line in lines):
        raise DocumentProcessingError(
            f"Aus '{path.name}' konnte kein Text extrahiert werden "
            "(gescanntes PDF? Dann bitte zuerst OCR anwenden)."
        )

    detected = _detect_document_meta(lines, path)
    doc_metadata: dict[str, Any] = {
        "doc_id": doc_id,
        "source": path.name,
        "source_path": str(path.resolve()),
        "file_type": suffix.lstrip("."),
        "n_pages": n_pages,
        "ingested_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        **detected,
    }
    if pdf_meta.get("title") and len(str(pdf_meta["title"]).strip()) > 5:
        doc_metadata["doc_title"] = str(pdf_meta["title"]).strip()
    if pdf_meta.get("author"):
        doc_metadata["author"] = str(pdf_meta["author"]).strip()

    units = _parse_units(lines)
    units = _merge_small_units(units, config.MIN_CHUNK_CHARS, chunk_size)
    chunks = units_to_chunks(units, doc_metadata, chunk_size, chunk_overlap)

    if not chunks:
        raise DocumentProcessingError(f"'{path.name}' ergab keine verwertbaren Chunks.")

    structured = sum(1 for c in chunks if c.metadata.get("paragraph") or c.metadata.get("annex"))
    if structured < len(chunks) * 0.2:
        warnings.append(
            "Nur wenige Chunks konnten einer Paragraphen-/Annex-Struktur "
            "zugeordnet werden - das Dokument folgt womoeglich nicht dem "
            "UNECE-Layout. Die Suche funktioniert trotzdem, die Zitate sind "
            "aber ungenauer."
        )

    doc_metadata["n_chunks"] = len(chunks)
    return ProcessedDocument(
        doc_id=doc_id,
        path=path,
        metadata=doc_metadata,
        chunks=chunks,
        warnings=warnings,
    )


def process_documents(paths: Iterable[str | Path]) -> list[ProcessedDocument]:
    """Verarbeitet mehrere Dateien; Fehler einzelner Dateien stoppen den Lauf nicht."""
    results: list[ProcessedDocument] = []
    for path in paths:
        try:
            results.append(process_document(path))
        except DocumentProcessingError as exc:
            logger.warning("Uebersprungen: %s", exc)
    return results


# --------------------------------------------------------------------------- #
# Hilfen fuer die UI (Vorschau & Metadaten)
# --------------------------------------------------------------------------- #
def get_document_info(path: str | Path) -> dict[str, Any]:
    """Liefert Metadaten fuer die Datei-Vorschau in der UI (ohne Chunking)."""
    path = Path(path)
    info: dict[str, Any] = {
        "filename": path.name,
        "size_mb": round(path.stat().st_size / (1024 * 1024), 2) if path.is_file() else 0.0,
        "file_type": path.suffix.lstrip(".").upper(),
        "pages": None,
        "title": "",
        "author": "",
        "regulation": "",
    }
    if not path.is_file():
        return info

    try:
        if path.suffix.lower() == ".pdf" and fitz is not None:
            with fitz.open(str(path)) as document:
                info["pages"] = document.page_count
                meta = document.metadata or {}
                info["title"] = (meta.get("title") or "").strip()
                info["author"] = (meta.get("author") or "").strip()
                head = "\n".join(
                    document.load_page(i).get_text("text")
                    for i in range(min(3, document.page_count))
                )
        elif path.suffix.lower() == ".docx" and docx is not None:
            document = docx.Document(str(path))
            paragraphs = [p.text for p in document.paragraphs[:200]]
            info["pages"] = None
            core = document.core_properties
            info["title"] = (core.title or "").strip()
            info["author"] = (core.author or "").strip()
            head = "\n".join(paragraphs)
        else:
            head = ""
    except Exception as exc:  # pragma: no cover
        logger.debug("Vorschau-Metadaten fehlgeschlagen: %s", exc)
        head = ""

    match = RE_REGULATION_NO.search(head) or RE_GTR_NO.search(head)
    if match:
        info["regulation"] = match.group(0).strip()
    elif RE_REG_IN_FILENAME.search(path.stem):
        info["regulation"] = f"UN R{RE_REG_IN_FILENAME.search(path.stem).group(1)}"
    if not info["title"]:
        info["title"] = path.stem
    return info


def render_pdf_page(path: str | Path, page_number: int = 1, zoom: float = 1.6) -> bytes | None:
    """Rendert eine PDF-Seite als PNG (fuer ``st.image``). ``page_number`` ist 1-basiert."""
    if fitz is None:
        return None
    path = Path(path)
    if not path.is_file() or path.suffix.lower() != ".pdf":
        return None
    try:
        with fitz.open(str(path)) as document:
            index = max(0, min(page_number - 1, document.page_count - 1))
            pixmap = document.load_page(index).get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            return pixmap.tobytes("png")
    except Exception as exc:  # pragma: no cover
        logger.debug("PDF-Rendering fehlgeschlagen: %s", exc)
        return None


__all__ = [
    "Chunk",
    "ProcessedDocument",
    "DocumentProcessingError",
    "process_document",
    "process_documents",
    "get_document_info",
    "render_pdf_page",
    "build_citation",
    "sanitize_metadata",
    "compute_doc_id",
]
